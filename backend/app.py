from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import uuid
from datetime import datetime
import json
import numpy as np
from sentence_transformers import SentenceTransformer
from pymongo import MongoClient
import PyPDF2
import docx
import requests
from io import BytesIO
import faiss
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
OLLAMA_BASE_URL = "http://localhost:11434"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

try:
    client = MongoClient('mongodb://localhost:27017/', serverSelectionTimeoutMS=5000)
    client.admin.command('ping')
    db = client['rag_bot']
    documents_collection = db['documents']
    conversations_collection = db['conversations']
    print("✅ MongoDB connected successfully")
except Exception as e:
    print(f"⚠️  MongoDB connection failed: {e}")
    print("⚠️  The app will work without persistence. Install MongoDB or start the service to enable document storage.")
    documents_collection = None
    conversations_collection = None
    client = None

embedding_model = SentenceTransformer('all-MiniLM-L6-v2')


faiss_index = None
document_chunks = []

class RAGPipeline:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.embedding_model = embedding_model
        
    def allowed_file(self, filename):
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def extract_text_from_file(self, file_path, filename):
        """Extract text from different file types"""
        text = ""
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        try:
            if file_extension == 'txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            elif file_extension == 'pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            elif file_extension == 'docx':
                doc = docx.Document(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting text from {filename}: {str(e)}")
            return None
            
        return text

    def process_document(self, file_path, filename):
        """Process a document and add it to the vector store"""
        global faiss_index, document_chunks
        
        # Extract text
        text = self.extract_text_from_file(file_path, filename)
        if not text:
            return False
            
        # Create document and split into chunks
        doc = Document(page_content=text, metadata={"source": filename})
        chunks = self.text_splitter.split_documents([doc])
        
        # Generate embeddings
        chunk_texts = [chunk.page_content for chunk in chunks]
        embeddings = self.embedding_model.encode(chunk_texts)
        
        # Store in MongoDB
        doc_id = str(uuid.uuid4())
        document_data = {
            "_id": doc_id,
            "filename": filename,
            "content": text,
            "chunks": [
                {
                    "text": chunk.page_content,
                    "metadata": chunk.metadata,
                    "embedding": embedding.tolist()
                }
                for chunk, embedding in zip(chunks, embeddings)
            ],
            "created_at": datetime.utcnow()
        }
        
        if documents_collection is not None:
            try:
                documents_collection.insert_one(document_data)
            except Exception as e:
                print(f"Warning: Could not store document in MongoDB: {e}")
        else:
            print("Warning: MongoDB not available - document not persisted")
        

        if faiss_index is None:
            dimension = embeddings.shape[1]
            faiss_index = faiss.IndexFlatL2(dimension)
            document_chunks = []
        
        faiss_index.add(embeddings.astype('float32'))
        
        for i, chunk in enumerate(chunks):
            document_chunks.append({
                "text": chunk.page_content,
                "metadata": chunk.metadata,
                "doc_id": doc_id
            })
        
        return True

    def similarity_search(self, query, k=5):
        """Search for similar chunks using FAISS"""
        global faiss_index, document_chunks
        
        if faiss_index is None or len(document_chunks) == 0:
            return []
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query])
        
        # Search in FAISS
        distances, indices = faiss_index.search(query_embedding.astype('float32'), k)
        
        # Return relevant chunks
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(document_chunks):
                chunk = document_chunks[idx]
                results.append({
                    "text": chunk["text"],
                    "metadata": chunk["metadata"],
                    "score": float(distances[0][i])
                })
        
        return results

    def generate_response(self, query, context_chunks, model="gemma3:1b"):
        """Generate response using Ollama"""
        context = "\n\n".join([chunk["text"] for chunk in context_chunks])
        
        prompt = f"""Based on the following context, please answer the question. If you cannot find the answer in the context, please say so.

Context:
{context}

Question: {query}

Answer:"""

        try:
            health_check = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=30)
            if health_check.status_code != 200:
                return "Error: Ollama is not responding. Please run 'ollama serve' in a terminal."
            
            models = health_check.json().get('models', [])
            model_names = [m.get('name', '') for m in models]
            allowed_models = ['gemma3:1b', 'mistral:7b-instruct-q2_K']
            
            if model not in model_names:
                # Try to find an available allowed model
                available_allowed_models = [m for m in allowed_models if m in model_names]
                if available_allowed_models:
                    model = available_allowed_models[0]
                    print(f"Using available model: {model}")
                else:
                    available = ', '.join(model_names) if model_names else 'none'
                    return f"Error: None of the allowed models ({allowed_models}) are available. Available models: {available}. Please install the required models."
            
            try:
                status_response = requests.post(f"{OLLAMA_BASE_URL}/api/show", json={"name": model}, timeout=15)
                if status_response.status_code != 200:
                    print(f"Warning: Could not verify model {model} status, continuing anyway")
            except:
                pass 
            
            response = requests.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.7,
                        "top_p": 0.9,
                        "num_predict": 2048
                    }
                },
                timeout=300 
            )
            
            if response.status_code == 200:
                return response.json().get("response", "Sorry, I couldn't generate a response.")
            else:
                error_msg = f"Ollama error (status {response.status_code})"
                try:
                    error_detail = response.json().get('error', '')
                    if error_detail:
                        error_msg += f": {error_detail}"
                except:
                    pass
                return f"Error: {error_msg}"
                
        except requests.exceptions.ConnectionError:
            return "Error: Cannot connect to Ollama. Please run 'ollama serve' in a terminal and make sure it's accessible at http://localhost:11434"
        except requests.exceptions.Timeout:
            return "Error: Ollama request timed out. The model might be loading or the request is taking too long. Please try again in a few moments. You can increase timeout in the code if needed."
        except Exception as e:
            return f"Error: {str(e)}"

# Initialize RAG pipeline
rag_pipeline = RAGPipeline()

# Load existing documents from MongoDB on startup
def load_existing_documents():
    global faiss_index, document_chunks
    
    if documents_collection is None:
        print("⚠️  MongoDB not available - no existing documents loaded")
        return
    
    try:
        docs = list(documents_collection.find())
        if not docs:
            print("ℹ️  No existing documents found in database")
            return
        
        all_embeddings = []
        document_chunks = []
        
        for doc in docs:
            for chunk in doc.get("chunks", []):
                embedding = np.array(chunk["embedding"])
                all_embeddings.append(embedding)
                document_chunks.append({
                    "text": chunk["text"],
                    "metadata": chunk["metadata"],
                    "doc_id": doc["_id"]
                })
        
        if all_embeddings:
            embeddings_array = np.vstack(all_embeddings)
            dimension = embeddings_array.shape[1]
            faiss_index = faiss.IndexFlatL2(dimension)
            faiss_index.add(embeddings_array.astype('float32'))
            print(f"✅ Loaded {len(docs)} documents with {len(document_chunks)} chunks")
    except Exception as e:
        print(f"⚠️  Error loading existing documents: {e}")

# API Routes
@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and rag_pipeline.allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        # Process the document
        success = rag_pipeline.process_document(file_path, filename)
        
        os.remove(file_path)
        
        if success:
            return jsonify({'message': f'File {filename} uploaded and processed successfully'})
        else:
            return jsonify({'error': 'Failed to process the document'}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/api/query', methods=['POST'])
def query():
    data = request.json
    user_query = data.get('query', '')
    model = data.get('model', 'gemma3:1b')
    conversation_id = data.get('conversation_id')
    
    # Validate that only allowed models are used
    allowed_models = ['gemma3:1b', 'mistral:7b-instruct-q2_K']
    if model not in allowed_models:
        return jsonify({'error': f'Model {model} is not allowed. Allowed models: {allowed_models}'}), 400
    
    if not user_query:
        return jsonify({'error': 'No query provided'}), 400
    
    # Search for relevant chunks
    relevant_chunks = rag_pipeline.similarity_search(user_query, k=5)
    
    if not relevant_chunks:
        return jsonify({
            'response': 'I don\'t have any relevant information to answer your question. Please upload some documents first.',
            'sources': [],
            'conversation_id': conversation_id
        })
    
    # Generate response
    response = rag_pipeline.generate_response(user_query, relevant_chunks, model)
    
    # Store conversation
    if not conversation_id:
        conversation_id = str(uuid.uuid4())
    
    conversation_data = {
        'conversation_id': conversation_id,
        'query': user_query,
        'response': response,
        'model': model,
        'sources': [chunk['metadata'].get('source', 'Unknown') for chunk in relevant_chunks],
        'timestamp': datetime.utcnow()
    }
    
    # Only store in MongoDB if available
    if conversations_collection is not None:
        try:
            conversations_collection.insert_one(conversation_data)
        except Exception as e:
            print(f"Warning: Could not store conversation in MongoDB: {e}")
    else:
        print("Warning: MongoDB not available - conversation not persisted")
    
    return jsonify({
        'response': response,
        'sources': list(set([chunk['metadata'].get('source', 'Unknown') for chunk in relevant_chunks])),
        'conversation_id': conversation_id
    })

@app.route('/api/models', methods=['GET'])
def get_models():
    """Get available Ollama models"""
    # Only allow these two specific models
    allowed_models = ['gemma3:1b', 'mistral:7b-instruct-q2_K']
    
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=15)
        if response.status_code == 200:
            models = response.json().get('models', [])
            model_names = [model['name'] for model in models]
            
            # Filter to only show the two allowed models that are actually available
            available_allowed_models = [m for m in allowed_models if m in model_names]
            
            if available_allowed_models:
                return jsonify({
                    'models': available_allowed_models,
                    'status': 'connected'
                })
            else:
                return jsonify({
                    'models': allowed_models,
                    'status': 'no_models',
                    'message': f'None of the allowed models found. Available models: {model_names}. Please install the required models.'
                })
        else:
            return jsonify({
                'models': allowed_models,
                'status': 'error',
                'message': f'Ollama responded with status {response.status_code}'
            })
    except requests.exceptions.ConnectionError:
        return jsonify({
            'models': allowed_models,
            'status': 'disconnected',
            'message': 'Cannot connect to Ollama. Run "ollama serve" in terminal.'
        })
    except Exception as e:
        return jsonify({
            'models': allowed_models,
            'status': 'error',
            'message': str(e)
        })

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get list of uploaded documents"""
    if documents_collection is None:
        return jsonify({
            'documents': [],
            'message': 'MongoDB not available - documents not persisted'
        })
    
    try:
        docs = list(documents_collection.find({}, {'filename': 1, 'created_at': 1}))
        return jsonify({
            'documents': [
                {
                    'id': str(doc['_id']),
                    'filename': doc['filename'],
                    'created_at': doc['created_at'].isoformat()
                }
                for doc in docs
            ]
        })
    except Exception as e:
        return jsonify({
            'documents': [],
            'error': f'Database error: {str(e)}'
        })

@app.route('/api/conversations', methods=['GET'])
def list_conversations():
    """List all conversation IDs with their first message and timestamp"""
    if conversations_collection is None:
        return jsonify({'conversations': [], 'message': 'MongoDB not available'}), 200
    try:
        pipeline = [
            {"$sort": {"timestamp": 1}},
            {"$group": {
                "_id": "$conversation_id",
                "first_query": {"$first": "$query"},
                "first_timestamp": {"$first": "$timestamp"},
                "model": {"$first": "$model"}
            }},
            {"$sort": {"first_timestamp": -1}}
        ]
        conversations = list(conversations_collection.aggregate(pipeline))
        return jsonify({
            'conversations': [
                {
                    'conversation_id': c['_id'],
                    'first_query': c.get('first_query', ''),
                    'first_timestamp': c.get('first_timestamp').isoformat() if c.get('first_timestamp') else '',
                    'model': c.get('model', '')
                }
                for c in conversations
            ]
        })
    except Exception as e:
        return jsonify({'conversations': [], 'error': str(e)}), 500

@app.route('/api/conversations/<conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    """Get conversation history"""
    conversations = list(conversations_collection.find(
        {'conversation_id': conversation_id},
        {'_id': 0}
    ).sort('timestamp', 1))
    
    return jsonify({'conversations': conversations})

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'message': 'RAG Bot API is running'})

if __name__ == '__main__':
    # Load existing documents on startup
    load_existing_documents()
    print("RAG Bot API starting...")
    if documents_collection is not None:
        try:
            doc_count = documents_collection.count_documents({})
            print(f"Documents in database: {doc_count}")
        except Exception as e:
            print(f"Could not count documents: {e}")
    else:
        print("MongoDB not available - documents will not be persisted")
    app.run(debug=True, port=5000)